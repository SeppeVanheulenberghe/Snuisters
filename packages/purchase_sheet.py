# -*- coding: utf-8 -*-
"""
Created on Fri Dec  9 22:47:20 2022.

@author: Seppe
"""

from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_LINE_SPACING
from docx.image.exceptions import UnrecognizedImageError
from typing import List
from packages.inventory import Inventory
from abc import ABC, abstractmethod
from papertrail_log.papertrail_logging import LoggingToPapertrail

log = LoggingToPapertrail()


class Snuisters_Document(ABC):
    """Responsible for writing the purchase sheets docx."""

    def __init__(self, inventory: Inventory):
        self.inventory = inventory
        self.inventory_AllItems = inventory.AllItems
        self.inventory_details = inventory.details

    def open_doc(self, template_name: str) -> None:
        """Initialize Document object."""
        self.doc = Document(template_name)

    def add_heading(self, header: str, level: int):
        """Add new heading to Document object"""
        return self.doc.add_heading(header, level)

    def add_paragraph(self):
        """Add new paragraph to Document object."""
        return self.doc.add_paragraph()

    def add_text_to_paragraph(self, paragraph_name: str, text: str):
        """Add text to already existing paragraph by name"""
        paragraph_name.add_run(text)

    def add_host_to_details_paragraph(self, details):
        """Add the host of the inventory to the Document object"""
        host = f"ONTVANGER: {self.inventory_details.host.name}"
        self.add_text_to_paragraph(details, host)

    def make_details_paragraph(self):
        self.add_heading("Details", 1)
        details = self.add_paragraph()
        self.add_host_to_details_paragraph(details)

    def make_doc_table(self, rows: int, cols: int) -> None:
        """Add table to document."""
        self.table = self.doc.add_table(rows, cols)

    def fill_table_header(self, labels: List[str]) -> None:
        """Fill in header labels in tabel."""
        header = self.table.rows[0].cells
        for i, label in enumerate(labels):
            header[i].text = label

    def fill_table(self) -> None:
        """Fill in table with inventory items."""
        for i in range(1, len(self.inventory) + 1):
            name = list(self.inventory_AllItems.keys())[i - 1]
            item = self.inventory_AllItems[name]

            row = self.table.rows[i].cells
            row[0].text = item.name
            paragraph = row[0].paragraphs[0]
            run = paragraph.add_run()
            try:
                run.add_picture(
                    "./images/" + self.inventory_AllItems[name].image, width=Cm(2)
                )

            except TypeError:
                row[0].text += "\n\nimage not found\n"
                log.error("TypeError: image not found")
            
            except UnrecognizedImageError:
                row[0].text += "\n\nimage not found\n"
                log.error("UnrecognizedImageError: image not found")

            row[1].text = "€ " + f"{item.price:.2f}"
            row[2].text = "€ " + f"{item.profit:.2f}"
            row[3].text = str(item.number_received)

    @property
    def get_inventory_details_host(self) -> str:
        """Get host of inventory from Details object."""
        return self.inventory_details.host

    def set_table_style(self, style: str) -> None:
        """Set the table style of the purchase sheet."""
        self.table.style = style

    @abstractmethod
    def create(
        self, doc_name: str, template_name: str, table_style: str = "Plain Table 1"
    ) -> None:
        """Create purchase sheet document."""


class Snuisters_Purchase_Sheet(Snuisters_Document):
    """Purchase sheet specifically for Snuisters."""

    header_labels: List[str] = [
        "ARTIKEL",
        "PRIJS",
        "WINSTMARGE",
        "AANTAL\nONTVANGEN",
        "AANTAL\nVERKOCHT",
    ]

    def create(
        self, doc_name: str, template_name: str, table_style: str = "Plain Table 1"
    ) -> None:
        """Create purchase sheet document."""
        # opening new document
        templates_dir = "templates/"
        self.open_doc(templates_dir + template_name + ".docx")
        # adding details about inventory
        self.make_details_paragraph()
        # making table for inventory items
        rows, cols = len(self.inventory) + 1, len(self.header_labels)
        self.make_doc_table(rows, cols)
        self.fill_table_header(self.header_labels)
        self.fill_table()
        self.set_table_style(table_style)
        self.doc.save("./documents/" + doc_name + ".docx")


class Snuisters_Invoice(Snuisters_Document):
    """Invoice document for Snuisters."""

    header_labels: List[str] = [
        "ARTIKEL",
        "PRIJS",
        "AANTAL\nVERKOCHT",
        "TOTALE\nVERKOOPPRIJS",
    ]

    def add_host_name_to_details_paragraph(self, details):
        """Add the host name of the inventory to the Document object"""
        host = f"{self.inventory_details.host.name}\n"
        self.add_text_to_paragraph(details, host)

    def add_host_company_name_to_details_paragraph(self, details):
        """Add the company of the host to details paragraph."""
        host = f"{self.inventory_details.host.company_name}\n"
        self.add_text_to_paragraph(details, host)

    def add_host_phone_number_to_details_paragraph(self, details):
        """Add the phone number of the host to details paragraph."""
        host = f"{self.inventory_details.host.phone_number}\n"
        self.add_text_to_paragraph(details, host)

    def add_host_address_street_to_details_paragraph(self, details):
        """Add the street of the host address to details paragraph."""
        host = f"{self.inventory_details.host.address.street}\n"
        self.add_text_to_paragraph(details, host)

    def add_host_address_city_to_details_paragraph(self, details):
        """Add the city of the host address to details paragraph."""
        host = f"{self.inventory_details.host.address.city}\n"
        self.add_text_to_paragraph(details, host)

    def make_details_paragraph(self):
        """Make new details paragraph specifically for an invoice document."""
        self.add_heading("AAN:", 1)
        details = self.add_paragraph()
        details.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        self.add_host_name_to_details_paragraph(details)
        self.add_host_company_name_to_details_paragraph(details)
        self.add_host_address_street_to_details_paragraph(details)
        self.add_host_address_city_to_details_paragraph(details)
        self.add_host_phone_number_to_details_paragraph(details)
        # self.add_host_to_details_paragraph(details)

    def fill_table(self) -> None:
        """Fill in table with inventory items."""
        for i in range(1, len(self.inventory) + 1):
            name = list(self.inventory_AllItems.keys())[i - 1]
            item = self.inventory_AllItems[name]

            row = self.table.rows[i].cells
            row[0].text = item.name
            paragraph = row[0].paragraphs[0]
            run = paragraph.add_run()
            run.add_break()
            try:
                run.add_picture(
                    "./images/" + self.inventory_AllItems[name].image, width=Cm(2)
                )
                run.add_break()

            except TypeError:
                row[0].text += "\n\nimage not found\n"
                log.error("TypeError: image not found")
            
            except UnrecognizedImageError:
                row[0].text += "\n\nimage not found\n"
                log.error("UnrecognizedImageError: image not found")


            row[1].text = "€ " + f"{item.price - item.profit:.2f}"
            row[2].text = f"{item.number_sold:.0f}"
            row[3].text = "€ " + f"{item.total_price:.2f}"

    def create(
        self, doc_name: str, template_name: str, table_style: str = "Plain Table 1"
    ) -> None:
        """Create invoice document."""
        # opening new document
        templates_dir = "templates/"
        self.open_doc(templates_dir + template_name + ".docx")
        # adding details about inventory
        self.make_details_paragraph()
        # make table
        # making table for inventory items
        rows, cols = len(self.inventory) + 1, len(self.header_labels)
        self.make_doc_table(rows, cols)
        self.fill_table_header(self.header_labels)
        self.fill_table()
        self.set_table_style(table_style)
        self.doc.save("./documents/" + doc_name + ".docx")
